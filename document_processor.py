"""document_processor.py — скачивание и извлечение текста из документов закупки."""

from __future__ import annotations

import hashlib
import logging
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Iterable
from urllib.parse import urljoin, urlparse

import requests
from bs4 import BeautifulSoup

import config
from scraper import HEADERS, BASE_URL

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".doc", ".docx", ".odt", ".pdf", ".txt", ".rtf", ".xlsx", ".xls", ".csv", ".zip"}
DOWNLOAD_MARKERS = (
    "/download/",
    "/filestore/",
    "file.html",
    "printform",
)
DOCUMENT_PAGE_MARKERS = (
    "documents.html",
)


def safe_filename(name: str, fallback: str = "document") -> str:
    name = re.sub(r"[\\/:*?\"<>|]+", "_", name).strip(" .")
    return name[:180] or fallback


def _repair_eis_href_entities(href: str) -> str:
    """Repair EIS query params decoded by HTML parsers as named entities."""
    return (
        (href or "")
        .replace("®Number", "&regNumber")
        .replace("¬iceGuid", "&noticeGuid")
    )


def find_document_items(page_html: str, base_url: str, *, include_pages: bool = True) -> list[tuple[str, str]]:
    soup = BeautifulSoup(page_html, "html.parser")
    items: list[tuple[str, str]] = []
    for a in soup.find_all("a", href=True):
        href = _repair_eis_href_entities(a.get("href", ""))
        raw_text = a.get_text(" ", strip=True)
        text = raw_text.lower()
        lower_href = href.lower()
        if "/rpt/" in lower_href or "zakupki-traffic" in lower_href:
            continue
        looks_like_file = _looks_like_document_file(lower_href, text)
        looks_like_page = include_pages and _looks_like_document_page(lower_href)
        looks_like_doc = looks_like_file or looks_like_page
        if not looks_like_doc:
            continue
        url = urljoin(base_url, href)
        if not any(existing == url for existing, _ in items):
            items.append((url, raw_text))
    return items[: config.MAX_DOCUMENTS_PER_TENDER]


def find_document_links(page_html: str, base_url: str, *, include_pages: bool = True) -> list[str]:
    return [url for url, _ in find_document_items(page_html, base_url, include_pages=include_pages)]


def _looks_like_document_file(lower_href: str, text: str = "") -> bool:
    return any(ext in lower_href for ext in SUPPORTED_EXTENSIONS) or any(
        marker in lower_href for marker in DOWNLOAD_MARKERS
    ) or any(word in text for word in ["извещение", "техническое", "проект договора", "скачать"])


def _looks_like_document_page(lower_href: str) -> bool:
    return any(marker in lower_href for marker in DOCUMENT_PAGE_MARKERS)


def _expand_document_items(items: list[tuple[str, str]]) -> list[tuple[str, str]]:
    expanded: list[tuple[str, str]] = []
    for link, label in items:
        lower = link.lower()
        if _looks_like_document_page(lower) and not _looks_like_document_file(lower):
            try:
                resp = requests.get(link, headers=HEADERS, timeout=30)
                if resp.status_code != 200 or not resp.text:
                    logger.info("Страница документов не скачана %s: HTTP %s", link, resp.status_code)
                    continue
                for nested, nested_label in find_document_items(resp.text, link, include_pages=False):
                    if not any(existing == nested for existing, _ in expanded):
                        expanded.append((nested, nested_label))
                        if len(expanded) >= config.MAX_DOCUMENTS_PER_TENDER:
                            return expanded
            except requests.RequestException as exc:
                logger.warning("Ошибка открытия страницы документов %s: %s", link, exc)
            continue
        if not any(existing == link for existing, _ in expanded):
            expanded.append((link, label))
            if len(expanded) >= config.MAX_DOCUMENTS_PER_TENDER:
                return expanded
    return expanded


def _filename_from_label(label: str) -> str:
    name = safe_filename(label or "")
    if Path(name).suffix.lower() in SUPPORTED_EXTENSIONS:
        return name
    return ""


def _guess_filename(response: requests.Response, url: str, index: int) -> str:
    cd = response.headers.get("content-disposition", "")
    match = re.search(r"filename\*=UTF-8''([^;]+)|filename=\"?([^\";]+)", cd, re.I)
    if match:
        raw = match.group(1) or match.group(2)
        try:
            from urllib.parse import unquote
            return safe_filename(unquote(raw))
        except Exception:
            return safe_filename(raw)
    path_name = Path(urlparse(url).path).name
    if path_name and "." in path_name:
        return safe_filename(path_name)
    content_type = response.headers.get("content-type", "").lower()
    ext = ".bin"
    if "pdf" in content_type:
        ext = ".pdf"
    elif "word" in content_type or "officedocument" in content_type:
        ext = ".docx"
    elif "zip" in content_type:
        ext = ".zip"
    elif "excel" in content_type or "spreadsheet" in content_type:
        ext = ".xlsx"
    elif "html" in content_type:
        ext = ".html"
    return f"document_{index}{ext}"


def download_documents(purchase_number: str, page_html: str, tender_url: str) -> dict:
    """Скачивает документы карточки закупки в data/documents/<purchase_number>."""
    target_dir = config.DOCUMENTS_DIR / safe_filename(purchase_number)
    target_dir.mkdir(parents=True, exist_ok=True)

    items = _expand_document_items(find_document_items(page_html, tender_url))
    saved: list[Path] = []
    for i, (link, label) in enumerate(items, start=1):
        if _looks_like_document_page(link.lower()) and not _looks_like_document_file(link.lower()):
            continue
        try:
            resp = requests.get(link, headers=HEADERS, timeout=30)
            if resp.status_code != 200 or not resp.content:
                logger.info("Документ не скачан %s: HTTP %s", link, resp.status_code)
                continue
            content_type = resp.headers.get("content-type", "").lower()
            suffix = Path(urlparse(link).path).suffix.lower()
            if "html" in content_type and suffix not in SUPPORTED_EXTENSIONS:
                logger.info("HTML-страница не считается документом %s", link)
                continue
            filename = _filename_from_label(label) or _guess_filename(resp, link, i)
            path = target_dir / filename
            path.write_bytes(resp.content)
            saved.append(path)
        except requests.RequestException as exc:
            logger.warning("Ошибка скачивания документа %s: %s", link, exc)

    return {"dir": str(target_dir), "files": saved, "links": [url for url, _ in items]}


def extract_text_from_file(path: Path) -> str:
    ext = path.suffix.lower()
    try:
        if ext == ".docx":
            return _extract_docx(path)
        if ext in {".doc", ".odt"}:
            return _extract_office_text(path)
        if ext == ".pdf":
            return _extract_pdf(path)
        if ext == ".xlsx":
            return _extract_xlsx(path)
        if ext == ".xls":
            return _extract_xls(path)
        if ext == ".zip":
            return _extract_zip(path)
        if ext in {".txt", ".csv", ".rtf", ".html", ".htm", ".xml"}:
            return _extract_plain(path)
        # Пробуем как текст, если расширение неизвестно.
        return _extract_plain(path)
    except Exception as exc:
        logger.warning("Не удалось извлечь текст из %s: %s", path.name, exc)
        return ""


def _extract_plain(path: Path) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "cp1251", "latin-1"):
        try:
            text = raw.decode(enc, errors="ignore")
            if path.suffix.lower() in {".html", ".htm", ".xml"}:
                soup = BeautifulSoup(text, "html.parser")
                return soup.get_text("\n", strip=True)
            return text
        except Exception:
            continue
    return ""


def _extract_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx не установлен, пропускаю %s", path.name)
        return ""
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _decode_command_output(raw: bytes) -> str:
    for enc in ("utf-8", "cp1251", "utf-16le", "latin-1"):
        try:
            text = raw.decode(enc, errors="ignore")
            if text.strip():
                return text
        except Exception:
            continue
    return ""


def _run_text_command(cmd: list[str], timeout: int = 45) -> str:
    try:
        result = subprocess.run(
            cmd,
            check=True,
            capture_output=True,
            timeout=timeout,
        )
    except Exception:
        return ""
    return _decode_command_output(result.stdout)


def _extract_with_libreoffice(path: Path) -> str:
    binary = shutil.which("soffice") or shutil.which("libreoffice")
    if not binary:
        return ""
    try:
        with tempfile.TemporaryDirectory() as tmp:
            outdir = Path(tmp)
            subprocess.run(
                [
                    binary,
                    "--headless",
                    "--convert-to",
                    "txt:Text",
                    "--outdir",
                    str(outdir),
                    str(path),
                ],
                check=True,
                capture_output=True,
                timeout=60,
            )
            out = outdir / f"{path.stem}.txt"
            if out.exists():
                return _decode_command_output(out.read_bytes())
    except Exception:
        return ""
    return ""


def _extract_office_text(path: Path) -> str:
    """Extract text from legacy Office/OpenDocument files on macOS and Linux.

    macOS ships `textutil`; Ubuntu deployments should install at least one of
    antiword, catdoc, libreoffice/soffice, or pandoc for old `.doc` files.
    """
    candidates: list[list[str]] = []
    if shutil.which("textutil"):
        candidates.append(["textutil", "-convert", "txt", "-stdout", str(path)])
    if shutil.which("antiword"):
        candidates.append(["antiword", str(path)])
    if shutil.which("catdoc"):
        candidates.append(["catdoc", str(path)])
    if shutil.which("pandoc"):
        candidates.append(["pandoc", "-t", "plain", str(path)])

    for cmd in candidates:
        text = _run_text_command(cmd)
        if text.strip():
            return text

    text = _extract_with_libreoffice(path)
    if text.strip():
        return text
    return _extract_plain(path)


def _extract_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf не установлен, пропускаю %s", path.name)
        return ""
    reader = PdfReader(str(path))
    pages = []
    for page in reader.pages[:30]:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def _extract_xlsx(path: Path) -> str:
    try:
        import openpyxl
    except ImportError:
        logger.warning("openpyxl не установлен, пропускаю %s", path.name)
        return ""
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets[:5]:
        parts.append(f"# Лист: {ws.title}")
        for row in ws.iter_rows(max_row=200, values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                parts.append(" | ".join(values))
    return "\n".join(parts)


def _extract_xls(path: Path) -> str:
    try:
        import xlrd
    except ImportError:
        logger.warning("xlrd не установлен, пробую LibreOffice для %s", path.name)
        text = _extract_with_libreoffice(path)
        return text if text.strip() else _extract_plain(path)

    try:
        wb = xlrd.open_workbook(str(path), on_demand=True)
    except Exception:
        text = _extract_with_libreoffice(path)
        return text if text.strip() else _extract_plain(path)

    parts = []
    try:
        for sheet_name in wb.sheet_names()[:5]:
            ws = wb.sheet_by_name(sheet_name)
            parts.append(f"# Лист: {sheet_name}")
            for row_idx in range(min(ws.nrows, 200)):
                values = []
                for value in ws.row_values(row_idx):
                    if value is None:
                        continue
                    cell = str(value).strip()
                    if cell:
                        values.append(cell)
                if values:
                    parts.append(" | ".join(values))
    finally:
        try:
            wb.release_resources()
        except Exception:
            pass
    return "\n".join(parts)


def _extract_zip(path: Path) -> str:
    parts = []
    extract_dir = path.parent / (path.stem + "_unzipped")
    extract_dir.mkdir(exist_ok=True)
    try:
        with zipfile.ZipFile(path) as zf:
            for member in zf.infolist()[:30]:
                if member.is_dir():
                    continue
                member_name = safe_filename(Path(member.filename).name)
                if not member_name:
                    continue
                out_path = extract_dir / member_name
                out_path.write_bytes(zf.read(member))
                if out_path.suffix.lower() in SUPPORTED_EXTENSIONS | {".html", ".htm", ".xml"}:
                    text = extract_text_from_file(out_path)
                    if text:
                        parts.append(f"\n# Файл из архива: {member.filename}\n{text}")
    except zipfile.BadZipFile:
        return ""
    return "\n".join(parts)


def collect_document_text(files: Iterable[Path], max_chars: int | None = None) -> str:
    limit = max_chars or config.MAX_DOCUMENT_TEXT_CHARS
    chunks = []
    total = 0
    for path in files:
        text = extract_text_from_file(Path(path))
        if not text.strip():
            continue
        clean = normalize_text(text)
        part = f"\n\n# Документ: {Path(path).name}\n{clean}"
        chunks.append(part)
        total += len(part)
        if total >= limit:
            break
    return "".join(chunks)[:limit]


def normalize_text(text: str) -> str:
    lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def hash_files(files: Iterable[Path]) -> str:
    h = hashlib.sha256()
    for path in sorted(Path(p) for p in files):
        if not path.exists() or not path.is_file():
            continue
        h.update(path.name.encode("utf-8", errors="ignore"))
        h.update(str(path.stat().st_size).encode())
    return h.hexdigest()


def extract_financial_terms(text: str) -> dict:
    """Эвристически достаёт обеспечение/аванс/оплату из текста страницы и документов."""
    result = {
        "application_security_amount": None,
        "contract_security_amount": None,
        "warranty_security_amount": None,
        "advance_percent": None,
        "payment_terms": "",
        "execution_days": None,
    }
    if not text:
        return result

    compact = re.sub(r"\s+", " ", text.lower())

    result["application_security_amount"] = _find_money_near(compact, ["обеспечение заявки", "размер обеспечения заявки"])
    result["contract_security_amount"] = _find_money_near(compact, ["обеспечение исполнения контракта", "обеспечение исполнения договора"])
    result["warranty_security_amount"] = _find_money_near(compact, ["обеспечение гарантийных обязательств"])

    advance_match = re.search(r"аванс[^%]{0,80}(\d{1,3})\s*%", compact)
    if advance_match:
        result["advance_percent"] = float(advance_match.group(1))

    payment_match = re.search(r"(оплата[^.]{0,240}(?:дн|рабоч|календар|акт|счет|счёт)[^.]{0,120})", compact)
    if payment_match:
        result["payment_terms"] = payment_match.group(1)[:500]

    days_match = re.search(r"(?:срок[^.]{0,80}|в течение\s+)(\d{1,3})\s*(?:рабочих|календарных)?\s*дн", compact)
    if days_match:
        result["execution_days"] = int(days_match.group(1))

    return result


# Маркеры содержательных требований к участнику (лежат в ИК/ТЗ/проекте контракта,
# а не на странице common-info, где только формулярные ссылки ст. 31 / ст. 14).
_REQUIREMENT_MARKERS: list[tuple[str, str, str]] = [
    ("experience",  r"опыт[а-я ]{0,25}(?:исполнен|поставк|оказан|выполнен|аналогич|сопоставим)",
     "Опыт исполнения аналогичных контрактов"),
    ("license",     r"наличи[ея][^.]{0,40}лиценз|лицензи[ия][^.]{0,40}(?:фстэк|фсб|деятельност)",
     "Лицензия (ФСТЭК/ФСБ/иная)"),
    ("sro",         r"\bсро\b|саморегулируем",
     "Членство в СРО"),
    ("add_req_31_2", r"ч\.?\s*2\s*ст\.?\s*31|дополнительны[ех]\s+требовани|постановлени[ея][^.]{0,40}2571|пп\s*рф\s*2571|№\s*2571",
     "Доптребования (ч. 2 ст. 31 / ПП РФ 2571)"),
    ("qualification", r"квалификац|квалифицированн|сертифицированн[ыйаяое]|сертификат\s+(?:соответстви|специалист)",
     "Квалификация / сертификаты"),
    ("staff",       r"штатн[аыо][^.]{0,30}(?:специалист|сотрудник|персонал)|наличи[ея][^.]{0,30}специалист",
     "Требования к персоналу"),
    ("bid_security", r"обеспечени[ея]\s+заявк",
     "Обеспечение заявки"),
]


def extract_participant_requirements(text: str) -> list[dict]:
    """
    Ищет в тексте документов (ИК/ТЗ/проект контракта) содержательные требования
    к участнику по маркерам. Возвращает список {type, label, snippet}.

    Это не юридически точный разбор, а сигнал «тут есть доптребование/опыт/лицензия» —
    чтобы requirements_special стал осмысленным (на странице common-info этого нет).
    """
    if not text:
        return []
    compact = re.sub(r"\s+", " ", text.lower())
    found: list[dict] = []
    seen: set[str] = set()
    for key, pattern, label in _REQUIREMENT_MARKERS:
        m = re.search(pattern, compact)
        if not m or key in seen:
            continue
        seen.add(key)
        i = m.start()
        snippet = compact[max(0, i - 50): i + 160].strip()
        found.append({"type": key, "label": label, "snippet": snippet})
    return found


def _clean_work_scope_line(line: str) -> str:
    line = re.sub(r"\s+", " ", line or "").strip(" \t\r\n;")
    line = re.sub(r"^[\-•–—*]\s*", "", line)
    line = re.sub(r"^\d+(?:\.\d+)*[.)]?\s*", "", line)
    return line.strip()


def _is_good_work_scope_line(line: str) -> bool:
    low = line.lower().replace("ё", "е")
    if len(line) < 18 or len(line) > 650:
        return False
    if re.fullmatch(r"[\d\s.,;:/№()%-]+", line):
        return False
    bad = (
        "содержание", "оглавление", "страница", "подпись", "заказчик:",
        "исполнитель:", "цена контракта", "цена договора", "начальная",
        "обоснование", "расчет нмц", "реквизиты", "извещение",
        "порядковый номер реестровой записи", "реестре российского программного",
        "перечень нормативных правовых актов", "приказ минфина",
        "перечень программного обеспечения, в отношении которого",
    )
    return not any(x in low for x in bad)


def _work_scope_window(text: str) -> str:
    compact = re.sub(r"[ \t]+", " ", re.sub(r"\r\n?", "\n", text or ""))
    low = compact.lower().replace("ё", "е")
    markers = [
        "перечень оказываемых услуг",
        "перечень выполняемых работ",
        "требования к оказанию услуг",
        "требования к выполнению работ",
        "описание объекта закупки",
        "техническое задание",
    ]
    positions = [low.find(marker) for marker in markers if low.find(marker) >= 0]
    start = min(positions) if positions else 0
    return compact[start:start + 12000]


def extract_work_scope(text: str, source: str = "") -> dict:
    """Извлекает человекочитаемое описание работ/услуг из ТЗ.

    Для сервисных закупок это отдельный блок от товарной спецификации: там часто
    нет количеств, зато есть перечень услуг и условия выполнения.
    """
    if not text:
        return {}

    window = _work_scope_window(text)
    lines = [_clean_work_scope_line(line) for line in window.split("\n")]
    lines = [line for line in lines if line]
    low_lines = [line.lower().replace("ё", "е") for line in lines]

    title = ""
    for line, low in zip(lines, low_lines):
        if (
            _is_good_work_scope_line(line)
            and any(x in low for x in ("оказание услуг", "выполнение работ", "сопровожд", "адаптац"))
        ):
            title = re.split(r",?\s*включает в себя:?", line, maxsplit=1, flags=re.I)[0].strip(" .;:")
            title = title[:300]
            break

    items: list[str] = []
    in_list = False
    for raw in window.split("\n"):
        line = _clean_work_scope_line(raw)
        low = line.lower().replace("ё", "е")
        if "включает в себя" in low or "перечень оказываемых услуг" in low or "перечень выполняемых работ" in low:
            in_list = True
            continue
        if in_list and re.match(r"^\s*(?:2(?:\.|\s)|порядок|требования к качеству|требования к безопасности)", raw, re.I):
            break
        if not in_list and not re.match(r"\s*(?:[-•–—*]|\d+[.)])\s+\S", raw):
            continue
        if (
            _is_good_work_scope_line(line)
            and any(x in low for x in ("услуг", "работ", "сопровожд", "адаптац", "разработк", "доработк"))
            and line not in items
        ):
            items.append(line[:500])
        if len(items) >= 10:
            break

    if not items:
        for line, low in zip(lines, low_lines):
            if (
                _is_good_work_scope_line(line)
                and any(x in low for x in ("оказание", "выполнение", "сопровожд", "адаптац", "услуг", "работ"))
                and line not in items
            ):
                items.append(line[:500])
            if len(items) >= 6:
                break

    conditions: list[str] = []
    for line, low in zip(lines, low_lines):
        if not _is_good_work_scope_line(line):
            continue
        if any(x in low for x in (
            "рабочие дни", "рабочем месте заказчика", "удаленное оказание",
            "выезд осуществляется", "без выноса информационных баз", "конфиденциаль",
        )) and line not in conditions:
            conditions.append(line[:650])
        if len(conditions) >= 5:
            break

    summary_candidates = [line for line in lines if _is_good_work_scope_line(line)]
    summary = " ".join(summary_candidates[:4])[:1000]
    result = {
        "title": title,
        "items": items,
        "conditions": conditions,
        "text": summary,
    }
    if source:
        result["source"] = source
    return {key: value for key, value in result.items() if value}


def _work_scope_doc_score(path: Path) -> int:
    name = _display_filename(path).lower().replace("ё", "е")
    score = 0
    if "техническ" in name:
        score += 80
    if "тз" in name or "техническое задание" in name:
        score += 80
    if "описание" in name and "объект" in name:
        score += 70
    if "услуг" in name or "работ" in name:
        score += 45
    if "прилож" in name:
        score += 20
    if any(x in name for x in ("контракт", "договор", "нмц", "расчет", "обоснован", "проект")):
        score -= 80
    return score


def _docx_block_items(doc):
    try:
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
    except ImportError:
        return

    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _work_scope_table_caption(prev_paragraphs: list[str]) -> str:
    for text in reversed(prev_paragraphs[-4:]):
        low = text.lower().replace("ё", "е")
        if (
            "таблица" in low
            or "список" in low
            or any(x in low for x in ("объекты предоставления услуг", "приоритет", "подменного оборудования"))
        ):
            return text.strip(" .:")
    return ""


def _looks_like_work_scope_table(rows: list[list[str]], caption: str) -> bool:
    if len(rows) < 2 or max((len(row) for row in rows), default=0) < 2:
        return False

    header_text = " ".join(rows[0])
    low = f"{caption} {header_text}".lower().replace("ё", "е")
    if "термин" in low and "определение" in low:
        return False
    if any(x in low for x in ("информационная карта", "обоснование начальной", "расчет нмц")):
        return False

    good_markers = (
        "объект", "услуг", "работ", "тип объекта", "количество",
        "уровень приоритета", "срок устранения", "неисправност",
        "подменного оборудования", "перечень",
    )
    return any(marker in low for marker in good_markers)


def _row_looks_like_table_header(row: list[str]) -> bool:
    if not row:
        return False
    first = row[0].strip()
    if re.fullmatch(r"\d{1,4}\.?", first):
        return False
    low = " ".join(row).lower().replace("ё", "е")
    return any(x in low for x in (
        "№", "номер", "наименование", "тип", "количество", "уровень",
        "срок", "объект", "услуг", "работ", "описание",
    ))


def _generic_table_columns(width: int) -> list[str]:
    if width == 2:
        return ["№", "Описание"]
    return [f"Колонка {i}" for i in range(1, width + 1)]


def _extract_work_scope_docx_tables(path: Path, max_tables: int = 6) -> list[dict]:
    if path.suffix.lower() != ".docx" or _work_scope_doc_score(path) <= 0:
        return []
    try:
        from docx import Document
    except ImportError:
        return []
    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.warning("Не удалось открыть DOCX для таблиц описания работ %s: %s", path.name, exc)
        return []

    tables: list[dict] = []
    prev_paragraphs: list[str] = []
    for block in _docx_block_items(doc):
        if hasattr(block, "text"):
            text = re.sub(r"\s+", " ", block.text or "").strip()
            if text:
                prev_paragraphs.append(text)
                prev_paragraphs = prev_paragraphs[-8:]
            continue

        rows = [_row_values(row) for row in block.rows]
        rows = [row for row in rows if any(row)]
        caption = _work_scope_table_caption(prev_paragraphs)
        if not _looks_like_work_scope_table(rows, caption):
            continue

        header_is_data = not _row_looks_like_table_header(rows[0])
        width = max(len(row) for row in rows)
        columns = _generic_table_columns(width) if header_is_data else rows[0]
        data_rows = rows if header_is_data else rows[1:]
        if not data_rows:
            continue
        tables.append({
            "caption": caption,
            "columns": columns,
            "rows": data_rows,
            "source": _display_filename(path),
        })
        if len(tables) >= max_tables:
            break
    return tables


def extract_work_scope_from_files(files: Iterable[Path | str], max_chars: int = 30000) -> dict:
    paths = [Path(path) for path in files or []]
    paths = [path for path in paths if path.is_file()]
    if not paths:
        return {}
    paths.sort(key=lambda path: (_work_scope_doc_score(path), -path.stat().st_size), reverse=True)

    chunks: list[str] = []
    tables: list[dict] = []
    source = ""
    total = 0
    for path in paths[:8]:
        if len(tables) < 6:
            tables.extend(_extract_work_scope_docx_tables(path, max_tables=6 - len(tables)))
        text = extract_text_from_file(path)
        if not text or len(text.strip()) < 80:
            continue
        chunks.append(text)
        total += len(text)
        if not source:
            source = _display_filename(path)
        if total >= max_chars:
            break

    if not chunks:
        return {}
    scope = extract_work_scope("\n\n".join(chunks), source=source)
    if tables:
        scope["tables"] = tables
    return scope


# ── Критерии оценки заявок («Как начисляются баллы») ─────────────────────────
# Эвристика без LLM. Для конкурса/котировок ищем стоимостные/нестоимостные
# критерии и их веса; для аукциона критериев нет — побеждает цена.

# Маркеры, по которым понимаем, что в тексте есть раздел оценки заявок.
_CRITERIA_MARKERS = [
    "критерии оценки", "критерий оценки", "порядок оценки", "значимость критери",
    "удельный вес", "коэффициент значимости", "стоимостные критери",
    "нестоимостные критери", "нестоимостной критери", "показатели оценки",
    "величина значимости",
]

# Известные ярлыки критериев → варианты их написания в документации.
_CRITERIA_LABELS: list[tuple[str, list[str]]] = [
    ("Цена контракта",            ["цена контракта", "цена договора", "стоимостн", "цена"]),
    ("Квалификация участника",    ["квалификаци"]),
    ("Опыт участника",            ["опыт"]),
    ("Деловая репутация",         ["деловая репутация", "репутаци"]),
    ("Обеспеченность ресурсами",  ["обеспеченность", "материально-техническ", "трудовыми ресурс"]),
    ("Качество",                  ["качество выполнения", "качественн"]),
]


def _windows_around(text: str, markers: list[str], window: int = 600,
                    max_total: int = 4000) -> list[str]:
    """Вырезает фрагменты текста вокруг маркеров (для дешёвой отправки в LLM).

    Возвращает список непустых уникальных окон; суммарно не длиннее max_total.
    """
    if not text:
        return []
    compact = re.sub(r"\s+", " ", text)
    low = compact.lower()
    spans: list[tuple[int, int]] = []
    for marker in markers:
        start = 0
        while True:
            idx = low.find(marker, start)
            if idx == -1:
                break
            spans.append((max(0, idx - window // 3), min(len(compact), idx + window)))
            start = idx + len(marker)
    if not spans:
        return []
    # Слить пересекающиеся окна
    spans.sort()
    merged: list[list[int]] = [list(spans[0])]
    for s, e in spans[1:]:
        if s <= merged[-1][1]:
            merged[-1][1] = max(merged[-1][1], e)
        else:
            merged.append([s, e])
    chunks: list[str] = []
    total = 0
    for s, e in merged:
        frag = compact[s:e].strip()
        if frag and frag not in chunks:
            chunks.append(frag)
            total += len(frag)
        if total >= max_total:
            break
    return chunks


def extract_evaluation_criteria(text: str) -> dict:
    """Эвристически определяет тип процедуры и критерии оценки заявок.

    Возвращает безопасную структуру даже на пустом/нерелевантном тексте.
    """
    result = {
        "procedure_hint": "unknown",
        "has_criteria": False,
        "criteria": [],
        "note": "",
    }
    if not text:
        result["note"] = "Текст документации не загружен — проверьте порядок оценки вручную."
        return result

    compact = re.sub(r"\s+", " ", text.lower())

    # 1. Тип процедуры
    if "электронный аукцион" in compact or "аукцион в электронной форме" in compact or "аукцион" in compact:
        procedure = "аукцион"
    elif "запрос котировок" in compact or "котировочн" in compact:
        procedure = "котировки"
    elif "конкурс" in compact:
        procedure = "конкурс"
    else:
        procedure = "unknown"
    result["procedure_hint"] = procedure

    # 2. Для аукциона критериев оценки нет — решает цена
    if procedure == "аукцион":
        result["note"] = (
            "Похоже на электронный аукцион: победитель определяется по цене, отдельные "
            "критерии оценки обычно не применяются. Но требования к участнику и заявке всё "
            "равно надо проверить."
        )
        return result

    # 3. Ищем критерии и веса
    has_markers = any(m in compact for m in _CRITERIA_MARKERS)
    criteria: list[dict] = []
    seen: set[str] = set()
    for label, variants in _CRITERIA_LABELS:
        for v in variants:
            idx = compact.find(v)
            if idx == -1:
                continue
            if label in seen:
                break
            seen.add(label)
            # Вес ищем ВПЕРЁД от ярлыка (иначе ловим % предыдущего критерия).
            fwd = compact[idx: idx + 160]
            wm = re.search(r"(\d{1,3})\s*%", fwd)
            weight = int(wm.group(1)) if wm and 0 < int(wm.group(1)) <= 100 else None
            snippet = compact[max(0, idx - 40): idx + 160]
            criteria.append({"label": label, "weight": weight, "snippet": snippet.strip()[:200]})
            break

    result["criteria"] = criteria
    result["has_criteria"] = bool(criteria) or has_markers

    if criteria:
        result["note"] = (
            "Заявки оценивают по критериям ниже. Новичку важно усилить нестоимостные "
            "(опыт, квалификация), а не только снижать цену."
        )
    elif procedure == "котировки":
        result["note"] = "Запрос котировок: обычно решает цена контракта."
    elif has_markers:
        result["note"] = (
            "В документации есть раздел оценки, но критерии не распознаны автоматически — "
            "проверьте раздел «Критерии оценки» вручную или нажмите «Разобрать подробнее»."
        )
    else:
        result["note"] = (
            "Критерии оценки в тексте не найдены — возможно, это аукцион (решает цена) "
            "или текст ТЗ ещё не загружен."
        )
    return result


# ── Спецификация позиций (MVP4b-1) ───────────────────────────────────────────
# Эвристика: строки с «количество + единица измерения» → товарные позиции.
# Для услуговых ТЗ (без количеств) возвращаем пустой список с пояснением.

SPEC_MARKERS = [
    "наименование", "кол-во", "количество", "ед. изм", "единица измерения",
    "технические характеристики", "спецификация", "перечень товаров", "товар",
]
_SPEC_SERVICE_MARKERS = [
    "оказание услуг", "выполнение работ", "сопровождение", "техническая поддержка",
    "администрирование", "обслуживание",
]
_SPEC_QTY_RE = re.compile(
    r"(\d+(?:[.,]\d+)?)\s*(шт|штук\w*|ед\.?|единиц\w*|компл\w*|комплект\w*|рабоч\w*\s+мест\w*)",
    re.IGNORECASE,
)


def _norm_unit(raw: str) -> str:
    r = raw.lower()
    if r.startswith("компл") or r.startswith("комплект"):
        return "компл"
    if r.startswith("услуг") or r.startswith("условн"):
        return "усл"
    if r.startswith("раб"):
        return "раб.место"
    if r.startswith("ед") or r.startswith("единиц"):
        return "ед"
    return "шт"


def _display_filename(path: Path) -> str:
    """Возвращает читаемое имя, если ЕИС-скачивание сохранило mojibake."""
    name = path.name
    try:
        return name.encode("latin1").decode("utf-8")
    except Exception:
        return name


def _looks_like_object_description(path: Path) -> bool:
    name = _display_filename(path).lower().replace("ё", "е")
    return "описание" in name and "объект" in name and "закуп" in name


def _extract_zip_document_files(path: Path) -> list[Path]:
    if path.suffix.lower() != ".zip":
        return []
    try:
        _extract_zip(path)
    except Exception:
        return []
    extract_dir = path.parent / (path.stem + "_unzipped")
    if not extract_dir.exists():
        return []
    return [
        item for item in extract_dir.rglob("*")
        if item.is_file() and item.suffix.lower() in {".docx", ".doc"}
    ]


def _looks_like_object_description_content(path: Path) -> bool:
    if _looks_like_object_description(path):
        return True
    try:
        text = extract_text_from_file(path)[:5000]
    except Exception:
        return False
    low = text.lower().replace("ё", "е")
    return (
        "описание объекта закупки" in low
        or "спецификация на поставку" in low
        or ("раздел 1. спецификация" in low and "предмет контракта" in low)
    )


def _row_values(row) -> list[str]:
    values: list[str] = []
    last = None
    for cell in row.cells:
        value = re.sub(r"\s+", " ", cell.text or "").strip()
        # python-docx repeats merged-cell values; collapse adjacent duplicates.
        if value and value != last:
            values.append(value)
        last = value
    return values


def _row_values_by_column(row) -> list[str]:
    return [re.sub(r"\s+", " ", cell.text or "").strip() for cell in row.cells]


def _parse_qty(raw: str) -> float | None:
    match = re.search(r"\d+(?:[,.]\d+)?", raw or "")
    if not match:
        return None
    try:
        return float(match.group(0).replace(",", "."))
    except ValueError:
        return None


def _is_object_item_row(values: list[str]) -> bool:
    if len(values) < 4:
        return False
    if not re.fullmatch(r"\d{1,4}\.?", values[0].strip()):
        return False
    if _parse_qty(values[-1]) is None:
        return False
    unit = values[-2].lower()
    return bool(re.search(r"штук|шт|единиц|компл|услуг|условн", unit))


def _find_header_index(values: list[str], *needles: str, exclude: str = "") -> int | None:
    for i, value in enumerate(values):
        low = value.lower().replace("ё", "е")
        if exclude and exclude in low:
            continue
        if all(needle in low for needle in needles):
            return i
    return None


def _find_any_header_index(values: list[str], variants: Iterable[tuple[str, ...]], exclude: str = "") -> int | None:
    for needles in variants:
        idx = _find_header_index(values, *needles, exclude=exclude)
        if idx is not None:
            return idx
    return None


def _append_requirement(item: dict, requirement: str) -> None:
    requirement = re.sub(r"\s+", " ", requirement or "").strip()
    if not requirement:
        return
    merged = item.get("requirements") or ""
    item["requirements"] = (merged + ("\n" if merged else "") + requirement)[:1500]


def _extract_indicator_table_items(table) -> list[dict]:
    """Парсит формат: позиция, наименование, ед. изм., количество, показатель, значение."""
    header: list[str] | None = None
    name_idx = unit_idx = qty_idx = indicator_idx = value_idx = None
    items_by_key: dict[tuple[int, str, float], dict] = {}

    for row in table.rows:
        values = _row_values_by_column(row)
        if not values:
            continue
        low = " ".join(values).lower().replace("ё", "е")
        if "наименование" in low and ("количество" in low or "кол-во" in low) and "единица" in low:
            header = values
            name_idx = _find_header_index(values, "наименование", exclude="показател")
            unit_idx = _find_header_index(values, "единица")
            qty_idx = _find_any_header_index(values, [("количество",), ("кол-во",)])
            indicator_idx = _find_any_header_index(values, [
                ("наименование", "показател"),
                ("наименование", "характерист"),
            ])
            value_idx = (
                _find_header_index(values, "содержание")
                or _find_header_index(values, "значение", "показател")
                or _find_header_index(values, "значение", "характерист")
                or _find_header_index(values, "значение")
            )
            continue
        if not header or name_idx is None or unit_idx is None or qty_idx is None:
            continue
        if len(values) <= max(name_idx, unit_idx, qty_idx):
            continue
        pos_match = re.fullmatch(r"(\d{1,4})\.?", values[0].strip())
        qty = _parse_qty(values[qty_idx])
        if not pos_match or qty is None:
            continue
        name = values[name_idx].strip()
        if not name:
            continue
        key = (int(pos_match.group(1)), name, qty)
        item = items_by_key.get(key)
        if not item:
            item = {
                "pos_no": key[0],
                "name": name[:200],
                "qty": qty,
                "unit": _norm_unit(values[unit_idx]),
                "requirements": "",
                "source": "object_description_docx",
            }
            items_by_key[key] = item
        if indicator_idx is not None and value_idx is not None and len(values) > max(indicator_idx, value_idx):
            _append_requirement(item, f"{values[indicator_idx]}: {values[value_idx]}")

    return sorted(items_by_key.values(), key=lambda item: item.get("pos_no") or 0)


def _is_unit_cell(value: str) -> bool:
    return bool(re.search(r"штук|шт|единиц|компл|услуг|условн", (value or "").lower()))


def _specific_item_name(cells: list[str]) -> str:
    for cell in cells:
        if re.search(r"\b(лицензи\w*|сертификат\w*|медиа-комплект|дистрибутив\w*|код активации)\b", cell, re.I):
            return cell[:200]
    base = " ".join(c for c in cells[:3] if c and not re.fullmatch(r"\d{1,4}", c))
    return base[:200] or "Позиция"


def _extract_legacy_doc_items(text: str) -> list[dict]:
    """Парсит старые .doc, где конвертер сохранил таблицу как ячейки через \\x07."""
    cells = [re.sub(r"\s+", " ", c).strip() for c in (text or "").split("\x07")]
    cells = [c for c in cells if c]
    starts: list[tuple[int, int, int, float]] = []
    for i in range(0, max(0, len(cells) - 4)):
        pos_match = re.fullmatch(r"(\d{1,4})\.?", cells[i])
        if not pos_match or re.fullmatch(r"\d{1,4}\.?", cells[i + 1] if i + 1 < len(cells) else ""):
            continue
        for j in range(i + 2, min(i + 18, len(cells) - 1)):
            qty = _parse_qty(cells[j + 1])
            if _is_unit_cell(cells[j]) and qty is not None:
                starts.append((i, j, int(pos_match.group(1)), qty))
                break

    items: list[dict] = []
    filler = re.compile(
        r"^(?:значение характеристики|участник закупки|качественная|количественная|"
        r"в соответствии с ктру|данные характеристики обусловлены)",
        re.I,
    )
    for idx, (start, unit_idx, pos_no, qty) in enumerate(starts):
        end = starts[idx + 1][0] if idx + 1 < len(starts) else len(cells)
        body = cells[start + 1:end]
        name = _specific_item_name(body)
        req_parts: list[str] = []
        for cell in body:
            if cell == name or filler.search(cell):
                continue
            if cell in {cells[unit_idx], cells[unit_idx + 1]}:
                continue
            req_parts.append(cell)
        items.append({
            "pos_no": pos_no,
            "name": name,
            "qty": qty,
            "unit": _norm_unit(cells[unit_idx]),
            "requirements": "; ".join(req_parts)[:1500],
            "source": "object_description_doc",
        })
    return items


def extract_object_description_items(files: Iterable[Path | str]) -> dict:
    """Извлекает позиции из типового файла ЕИС «Описание объекта закупки».

    В таких DOCX количество часто лежит в отдельных табличных колонках
    «Единица измерения» / «Кол-во», поэтому текстовая эвристика `5 шт`
    их не видит.
    """
    paths = [Path(p) for p in files or []]
    expanded_paths: list[Path] = []
    for path in paths:
        expanded_paths.append(path)
        if path.suffix.lower() == ".zip":
            expanded_paths.extend(_extract_zip_document_files(path))
    paths = expanded_paths
    candidates = [
        p for p in paths
        if p.suffix.lower() in {".docx", ".doc"} and _looks_like_object_description_content(p)
    ]
    if not candidates:
        return {"items": [], "note": "Файл «Описание объекта закупки» не найден."}

    items: list[dict] = []
    for path in candidates:
        if path.suffix.lower() == ".doc":
            legacy_items = _extract_legacy_doc_items(_extract_office_text(path))
            if legacy_items:
                return {"items": legacy_items, "note": f"Извлечено из {_display_filename(path)}."}
            continue

        try:
            from docx import Document
        except ImportError:
            return {"items": [], "note": "python-docx не установлен."}
        try:
            doc = Document(str(path))
        except Exception as exc:
            logger.warning("Не удалось открыть описание объекта закупки %s: %s", path.name, exc)
            continue

        for table in doc.tables:
            indicator_items = _extract_indicator_table_items(table)
            if indicator_items:
                return {"items": indicator_items, "note": f"Извлечено из {_display_filename(path)}."}

        current: dict | None = None
        for table in doc.tables:
            for row in table.rows:
                values = _row_values(row)
                if not values:
                    continue
                if _is_object_item_row(values):
                    if current:
                        items.append(current)
                    pos_match = re.fullmatch(r"(\d{1,4})\.?", values[0].strip())
                    current = {
                        "pos_no": int(pos_match.group(1)) if pos_match else len(items) + 1,
                        "name": values[1][:200],
                        "qty": _parse_qty(values[-1]),
                        "unit": _norm_unit(values[-2]),
                        "requirements": "",
                        "source": "object_description_docx",
                    }
                    continue
                if current and "значение характеристики:" in " ".join(values).lower():
                    req = re.sub(r"(?i)^значение характеристики:\s*", "", values[-1]).strip()
                    _append_requirement(current, req)
        if current:
            items.append(current)
        if items:
            return {"items": items, "note": f"Извлечено из {_display_filename(path)}."}

    return {"items": [], "note": "В описании объекта закупки позиции не распознаны."}


def extract_spec_items(text: str) -> dict:
    """Эвристически извлекает товарные позиции из ТЗ. Возвращает {items, note}."""
    if not text:
        return {"items": [], "note": "Текст документации не загружен."}

    low = text.lower()
    qty_found = _SPEC_QTY_RE.search(text) is not None
    service_hits = sum(1 for m in _SPEC_SERVICE_MARKERS if m in low)
    if service_hits >= 2 and not qty_found:
        return {"items": [], "note": "Позиции поставки не найдены — похоже на услугу/работы, "
                                     "а не поставку оборудования."}

    items: list[dict] = []
    pos = 0
    for seg in re.split(r"[\n;]+", text):
        seg = seg.strip()
        if not seg:
            continue
        m = _SPEC_QTY_RE.search(seg)
        if not m:
            continue
        try:
            qty = float(m.group(1).replace(",", "."))
        except ValueError:
            continue
        unit = _norm_unit(m.group(2))
        name = re.sub(r"^\s*\d+[\.\)]\s*", "", seg[:m.start()]).strip(" .—-:\t")
        requirements = seg[m.end():].strip(" .,;—-")
        if not name:
            name = (requirements[:80] or "Позиция").strip()
            requirements = ""
        pos += 1
        items.append({
            "pos_no": pos, "name": name[:200], "qty": qty, "unit": unit,
            "requirements": requirements[:500], "source": "heuristic",
        })

    note = "" if items else ("Позиции с количеством не найдены — добавьте вручную "
                             "или попробуйте «Извлечь точнее (LLM)».")
    return {"items": items, "note": note}


def _find_money_near(text: str, markers: list[str]) -> float | None:
    for marker in markers:
        idx = text.find(marker)
        if idx == -1:
            continue
        fragment = text[idx : idx + 600]
        matches = re.findall(r"(\d[\d\s]{1,15}(?:[,.]\d{1,2})?)\s*(?:руб|₽)", fragment)
        for raw in matches:
            value = _parse_money(raw)
            if value and value > 0:
                return value
    return None


def _parse_money(raw: str) -> float | None:
    cleaned = raw.replace(" ", "").replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None
