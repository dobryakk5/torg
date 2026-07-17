import decision_aid
import config
import database
import document_processor
import filter_engine
import scraper


def test_dom_parser_extracts_stage_security_and_support():
    html = """
    <div class="blockInfo">
      <h2 class="blockInfo__title">Общая информация о закупке</h2>
      <section class="blockInfo__section">
        <span class="section__title">Этап закупки</span>
        <span class="section__info">Подача заявок</span>
      </section>
    </div>
    <div class="blockInfo">
      <h2 class="blockInfo__title">Обеспечение исполнения контракта</h2>
      <section class="blockInfo__section">
        <span class="section__title">Требуется обеспечение исполнения контракта</span>
        <span class="section__info">Да</span>
      </section>
      <section class="blockInfo__section">
        <span class="section__title">Размер обеспечения исполнения контракта</span>
        <span class="section__info">5 %</span>
      </section>
    </div>
    <div class="blockInfo">
      <h2 class="blockInfo__title">Информация о банковском и (или) казначейском сопровождении контракта</h2>
      <section class="blockInfo__section">
        <span class="section__info">Банковское или казначейское сопровождение контракта не требуется</span>
      </section>
    </div>
    """

    details = scraper.parse_common_info_details_from_html(html, price=373104)

    assert details["common"]["purchase_stage"] == "Подача заявок"
    assert details["contract_security"]["required"] is True
    assert details["contract_security"]["raw_value"] == "5 %"
    assert details["contract_security"]["percent"] == 5
    assert details["contract_security"]["amount"] == 18655.2
    assert details["contract_support"]["required"] is False


def test_treasury_support_not_required_does_not_stop_f8():
    text = filter_engine._normalize(
        "Информация о банковском и казначейском сопровождении контракта. "
        "Банковское или казначейское сопровождение контракта не требуется."
    )

    score = filter_engine._filter_contract_risks({}, text, "stage2")

    assert score.stop_factor is False
    assert not any("казначейское сопровождение" in signal for signal in score.signals)


def test_treasury_support_required_still_stops_f8():
    text = filter_engine._normalize("Казначейское сопровождение контракта требуется.")

    score = filter_engine._filter_contract_risks({}, text, "stage2")

    assert score.stop_factor is True
    assert any("казначейское сопровождение" in signal for signal in score.signals)


def test_detail_rejected_is_not_closed_status():
    assert decision_aid._gate_status({"status": "detail_rejected"}) is None


def test_finance_gate_reads_contract_security_from_details_json():
    gate = decision_aid._gate_finance_load({
        "price": 373104,
        "contract_security_amount": None,
        "details_json": {"contract_security": {"amount": 18655.2, "percent": 5}},
    })

    assert gate["status"] == "ok"
    assert "18 655" in gate["explain"]


def test_deadline_is_expired_for_old_procurement():
    now = database.datetime(2026, 6, 12, 12, 0, tzinfo=database.timezone.utc)

    assert database.deadline_is_expired("13.02.2014", now=now) is True


def test_date_only_deadline_stays_active_until_end_of_day():
    noon = database.datetime(2026, 6, 12, 12, 0, tzinfo=database.timezone.utc)
    next_day = database.datetime(2026, 6, 13, 0, 1, tzinfo=database.timezone.utc)

    assert database.deadline_is_expired("12.06.2026", now=noon) is False
    assert database.deadline_is_expired("12.06.2026", now=next_day) is True


def test_reload_llm_secrets_reads_updated_dotenv(tmp_path, monkeypatch):
    (tmp_path / ".env").write_text(
        "OPENROUTER_API_KEY=new-key\nOPENROUTER_BASE_URL=https://llm.example/v1\n",
        encoding="utf-8",
    )
    monkeypatch.setattr(config, "BASE_DIR", tmp_path)
    monkeypatch.setattr(config, "OPENROUTER_API_KEY", "")

    config.reload_llm_secrets()

    assert config.OPENROUTER_API_KEY == "new-key"
    assert config.OPENROUTER_BASE_URL == "https://llm.example/v1"


def test_criteria_api_uses_full_text_fallback_and_refreshes_env(monkeypatch):
    import asyncio
    import json
    import llm_analyzer
    import llm_provider
    import web_app

    refreshed = []
    captured = []
    monkeypatch.setattr(web_app.db, "get_tender", lambda pnum: {"purchase_number": pnum})
    monkeypatch.setattr(web_app, "_decide_text", lambda tender: "Электронный аукцион. Побеждает цена.")
    monkeypatch.setattr(web_app.config, "reload_llm_secrets", lambda: refreshed.append(True))
    monkeypatch.setattr(llm_provider, "is_configured", lambda: bool(refreshed))
    monkeypatch.setattr(
        llm_analyzer,
        "extract_criteria_llm",
        lambda chunks: captured.extend(chunks) or {"summary": "Разбор", "criteria": []},
    )

    response = asyncio.run(web_app.api_decide_criteria("lot-1"))
    body = json.loads(response.body)

    assert refreshed == [True]
    assert captured == ["Электронный аукцион. Побеждает цена."]
    assert body["llm"]["summary"] == "Разбор"


def test_print_form_is_not_treated_as_document():
    html = """
    <a href="/epz/order/notice/printForm/listModal.html?regNumber=0842600001626000024">
      Печатная форма
    </a>
    <a href="/filestore/document.docx">Техническое задание</a>
    """

    links = document_processor.find_document_links(html, "https://zakupki.gov.ru/")

    assert links == ["https://zakupki.gov.ru/filestore/document.docx"]


def test_eis_44_documents_use_notice_view_page():
    url = (
        "https://zakupki.gov.ru/epz/order/notice/printForm/listModal.html?"
        "regNumber=0842600001626000024"
    )

    assert scraper.to_documents_url(url) == (
        "https://zakupki.gov.ru/epz/order/notice/ea20/view/documents.html?"
        "regNumber=0842600001626000024"
    )


def test_eis_44_documents_keep_notice_type_from_url():
    url = (
        "https://zakupki.gov.ru/epz/order/notice/zk20/view/common-info.html?"
        "regNumber=0318500001226000005"
    )

    assert scraper.to_documents_url(url) == (
        "https://zakupki.gov.ru/epz/order/notice/zk20/view/documents.html?"
        "regNumber=0318500001226000005"
    )


def test_eis_223_documents_use_notice223_page():
    url = (
        "https://zakupki.gov.ru/epz/order/notice/notice223/common-info.html?"
        "noticeInfoId=123&regNumber=32211111111"
    )

    assert scraper.to_documents_url(url) == (
        "https://zakupki.gov.ru/epz/order/notice/notice223/documents.html?"
        "noticeInfoId=123&purchaseNoticeNumber=32211111111"
    )


def test_eis_223_documents_include_notice_guid():
    guid = "34aac1c7-dbad-4b90-b11e-a93bcd769eff"
    url = "https://zakupki.gov.ru/epz/order/notice/notice223/common-info.html?regNumber=32616194818"

    assert scraper.extract_notice_guid(
        f'<a href="printForm/view.html?purchaseNoticeGuid={guid}&amp;purchaseNoticeNumber=32616194818">'
    ) == guid
    assert scraper.to_documents_url(url, guid) == (
        "https://zakupki.gov.ru/epz/order/notice/notice223/documents.html?"
        f"purchaseNoticeNumber=32616194818&noticeGuid={guid}"
    )


def test_web_documents_url_includes_notice_guid():
    import web_app

    guid = "34aac1c7-dbad-4b90-b11e-a93bcd769eff"
    url = "https://zakupki.gov.ru/epz/order/notice/notice223/common-info.html?regNumber=32616194818"

    assert web_app.zakupki_documents_url(url, "32616194818", guid) == (
        "https://zakupki.gov.ru/epz/order/notice/notice223/documents.html?"
        f"purchaseNoticeNumber=32616194818&noticeGuid={guid}"
    )


def test_web_document_fetch_uses_existing_files_without_network(tmp_path, monkeypatch):
    import web_app

    document = tmp_path / "Техническое задание.txt"
    document.write_text("Локально сохранённый документ", encoding="utf-8")
    monkeypatch.setattr(
        scraper,
        "get_tender_page",
        lambda *args, **kwargs: (_ for _ in ()).throw(AssertionError("network must not be used")),
    )
    stages = []

    result = web_app._fetch_tender_documents(
        {
            "purchase_number": "32699999999",
            "documents_dir": str(tmp_path),
            "url": "https://zakupki.gov.ru/",
        },
        progress_cb=lambda stage, **data: stages.append(stage),
    )

    assert result["files"] == [document]
    assert "Локально сохранённый документ" in result["text"]
    assert stages == ["cached"]


def test_web_documents_dir_prefers_current_project_cache(tmp_path, monkeypatch):
    import web_app

    current_dir = tmp_path / "32616194818"
    current_dir.mkdir()
    (current_dir / "Техническое задание.txt").write_text("Документ", encoding="utf-8")
    monkeypatch.setattr(web_app.config, "DOCUMENTS_DIR", tmp_path)

    result = web_app._tender_documents_dir({
        "purchase_number": "32616194818",
        "documents_dir": "/path/from/another/server/32616194818",
    })

    assert result == current_dir


def test_tender_detail_opens_expired_tender_from_db(monkeypatch):
    import asyncio
    import web_app

    tender = {
        "purchase_number": "32616106894",
        "title": "Архивный тендер",
        "deadline": "13.02.2014",
        "price": 1000,
    }

    monkeypatch.setattr(web_app.db, "get_tender", lambda purchase_number: tender)
    monkeypatch.setattr(web_app.db, "list_tender_items", lambda purchase_number: [])
    monkeypatch.setattr(web_app.db, "get_tender_details", lambda purchase_number: {})
    monkeypatch.setattr(decision_aid, "build_card", lambda tender, text=None: {"verdict": "view"})
    monkeypatch.setattr(document_processor, "extract_evaluation_criteria", lambda text: {})
    monkeypatch.setattr(web_app, "_decide_text", lambda tender: "")
    monkeypatch.setattr(
        web_app.templates,
        "TemplateResponse",
        lambda request, template_name, context: {"template": template_name, "context": context},
    )

    response = asyncio.run(web_app.tender_detail(None, "32616106894"))

    assert response["template"] == "detail.html"
    assert response["context"]["tender"]["purchase_number"] == "32616106894"


def test_work_scope_preserves_docx_tables(tmp_path):
    from docx import Document

    path = tmp_path / "Описание объекта закупки.docx"
    doc = Document()
    doc.add_paragraph("Описание объекта закупки")
    doc.add_paragraph("Оказание услуг по сопровождению системы")
    doc.add_paragraph("Таблица 1 – Объекты предоставления услуг:")
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "№ п/п"
    table.rows[0].cells[1].text = "Тип объекта"
    table.rows[0].cells[2].text = "Количество"
    for idx in range(1, 13):
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = f"Коммутатор {idx}"
        row[2].text = "1 шт."
    doc.add_paragraph("Используемые термины и сокращения")
    terms = doc.add_table(rows=2, cols=2)
    terms.rows[0].cells[0].text = "Термин"
    terms.rows[0].cells[1].text = "Определение"
    terms.rows[1].cells[0].text = "Заявка"
    terms.rows[1].cells[1].text = "Обращение Заказчика"
    doc.save(path)

    scope = document_processor.extract_work_scope_from_files([path])

    assert len(scope["tables"]) == 1
    assert scope["tables"][0]["columns"] == ["№ п/п", "Тип объекта", "Количество"]
    assert len(scope["tables"][0]["rows"]) == 12
    assert scope["tables"][0]["rows"][-1] == ["12", "Коммутатор 12", "1 шт."]
